#!/usr/bin/env python3
"""Collection-sheet generator (step.md step 3r.6a, plan.md §16
"Collecting real handwriting samples"). A .docx table, one row per digit
0-9, with a configurable number of empty cells per row for handwritten
samples. The row position *is* the label — no manual annotation step,
the same self-labelling principle behind harvest.py's filename-encoded
labels.

Reuses a fix already applied once to marks-grid-template.docx (learn.md
step 0): python-docx's `row.height` is silently ignored by Word unless
`row.height_rule` is also set explicitly — easy to miss since python-docx
doesn't raise or warn, the row just quietly comes out the wrong size.

    python generate_collection_sheet.py --out collection_sheet.docx
    python generate_collection_sheet.py --samples-per-digit 15 --out sheet.docx

Collect from at least 4 different writers for ID/serial training data —
the instructor's own handwriting is nearly useless for that field, since
students write it, not the instructor (plan.md §16). Collect the
instructor's own handwriting separately, on its own sheet, for marks
training data instead.

This generates the *blank* sheet only. Turning a photographed, filled-in
sheet into `training_data/<writer>/<digit>/<uuid>.png` crops is step
3r.6a's other half, deliberately not built yet — same reasoning as
`detect.py` itself (step 1): tuning a detector against a table shape
nobody has actually photographed yet means guessing, not measuring. That
script gets built and tuned against a real filled sheet once one exists,
not before.
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

DIGITS = 10
SAMPLES_PER_DIGIT = 20
ROW_HEIGHT = Cm(1.8)  # tall enough to write one digit by hand comfortably —
                        # same order of magnitude as the marks-grid
                        # template's own answer-row fix (learn.md step 0)
LABEL_COL_WIDTH = Cm(1.2)
SAMPLE_COL_WIDTH = Cm(1.2)


def build_document(samples_per_digit: int) -> Document:
    doc = Document()
    doc.add_heading("Digit Handwriting Collection Sheet", level=1)
    doc.add_paragraph(
        "Write one digit per box, at normal handwriting speed — not "
        "carefully formed. Include genuinely messy variants; the point is "
        "covering the cases that fail, not the ones that already work."
    )

    table = doc.add_table(rows=DIGITS, cols=samples_per_digit + 1)
    table.style = "Table Grid"

    for digit in range(DIGITS):
        row = table.rows[digit]
        row.height = ROW_HEIGHT
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY  # see module docstring

        label_cell = row.cells[0]
        label_cell.width = LABEL_COL_WIDTH
        paragraph = label_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(digit))
        run.bold = True
        run.font.size = Pt(14)

        for col in range(1, samples_per_digit + 1):
            row.cells[col].width = SAMPLE_COL_WIDTH

    return doc


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--samples-per-digit", type=int, default=SAMPLES_PER_DIGIT)
    parser.add_argument("--out", type=Path, default=Path("collection_sheet.docx"))
    args = parser.parse_args()

    doc = build_document(args.samples_per_digit)
    doc.save(str(args.out))
    total = args.samples_per_digit * DIGITS
    print(f"wrote {args.out} — {args.samples_per_digit} samples x {DIGITS} digits = {total} cells")
    return 0


if __name__ == "__main__":
    sys.exit(main())
